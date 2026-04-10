from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List
import html
import json
import re

import requests
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from agents.api_handlers import ChatMessage, MultiProviderClients
from agents.config import settings


@dataclass(frozen=True)
class DraftResult:
    kind: str
    path: str
    preview: str


class SkillToolkit:
    def __init__(self, output_dir: str | None = None) -> None:
        self.output_dir = Path(output_dir or settings.artifacts_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.clients = MultiProviderClients()

    def _slugify(self, name: str) -> str:
        value = re.sub(r"[^a-zA-Z0-9]+", "-", name.strip().lower()).strip("-")
        return value or "artifact"

    def generate_document(
        self, title: str, markdown: str, output_format: str = "docx"
    ) -> Dict[str, Any]:
        slug = self._slugify(title)
        if output_format.lower() == "docx":
            path = self.output_dir / f"{slug}.docx"
            document = Document()
            document.add_heading(title, level=1)
            for line in markdown.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("# "):
                    document.add_heading(stripped[2:].strip(), level=2)
                elif stripped.startswith("## "):
                    document.add_heading(stripped[3:].strip(), level=3)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    document.add_paragraph(stripped[2:].strip(), style="List Bullet")
                else:
                    document.add_paragraph(stripped)
            document.save(str(path))
        else:
            path = self.output_dir / f"{slug}.pdf"
            document = SimpleDocTemplate(str(path), pagesize=letter)
            styles = getSampleStyleSheet()
            story: List[Any] = [
                Paragraph(html.escape(title), styles["Title"]),
                Spacer(1, 12),
            ]
            for line in markdown.splitlines():
                stripped = line.strip()
                if not stripped:
                    story.append(Spacer(1, 6))
                    continue
                if stripped.startswith("# "):
                    story.append(
                        Paragraph(html.escape(stripped[2:].strip()), styles["Heading1"])
                    )
                elif stripped.startswith("## "):
                    story.append(
                        Paragraph(html.escape(stripped[3:].strip()), styles["Heading2"])
                    )
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    story.append(
                        Paragraph(
                            "&bull; " + html.escape(stripped[2:].strip()),
                            styles["BodyText"],
                        )
                    )
                else:
                    story.append(Paragraph(html.escape(stripped), styles["BodyText"]))
            document.build(story)

        return {
            "kind": output_format.lower(),
            "path": str(path),
            "preview": markdown[:1000],
        }

    def _refine_image_prompt(self, prompt: str) -> str:
        messages = [
            ChatMessage(
                role="system",
                content=(
                    "You are an image prompt engineer. Rewrite the user's idea into one concise, vivid prompt "
                    "for image generation. Return only the prompt text."
                ),
            ),
            ChatMessage(role="user", content=prompt),
        ]
        try:
            refined = self.clients.groq_chat(
                messages, model=settings.groq_model_strong, temperature=0.2
            )
            return refined.strip() or prompt
        except Exception:  # noqa: BLE001
            return prompt

    def illustrate(self, prompt: str, provider: str = "groq") -> Dict[str, Any]:
        if provider not in {"groq", "huggingface"}:
            raise ValueError("Illustration provider must be 'groq' or 'huggingface'.")

        refined_prompt = (
            self._refine_image_prompt(prompt) if provider == "groq" else prompt
        )

        if not settings.huggingface_api_token:
            fallback_path = (
                self.output_dir / f"{self._slugify(prompt)[:50]}-image-prompt.txt"
            )
            fallback_path.write_text(refined_prompt, encoding="utf-8")
            return {
                "status": "fallback",
                "provider": provider,
                "path": str(fallback_path),
                "note": "Hugging Face image generation is unavailable, so the refined image prompt was saved instead.",
                "prompt": refined_prompt,
            }

        model = settings.huggingface_image_model
        url = f"https://api-inference.huggingface.co/models/{model}"
        response = requests.post(
            url,
            headers={"Authorization": f"Bearer {settings.huggingface_api_token}"},
            json={"inputs": refined_prompt},
            timeout=180,
        )
        if response.status_code >= 400:
            fallback_path = (
                self.output_dir / f"{self._slugify(prompt)[:50]}-image-prompt.txt"
            )
            fallback_path.write_text(refined_prompt, encoding="utf-8")
            return {
                "status": "fallback",
                "provider": provider,
                "path": str(fallback_path),
                "detail": response.text[:1000],
                "model": model,
                "note": "Image generation failed, so the refined image prompt was saved instead.",
                "prompt": refined_prompt,
            }

        # HF inference for image models usually returns binary image bytes.
        image_path = (
            self.output_dir
            / f"{self._slugify(prompt)[:50]}-{model.replace('/', '-')}.png"
        )
        image_path.write_bytes(response.content)
        return {
            "status": "ok",
            "provider": provider,
            "model": model,
            "path": str(image_path),
            "note": "Groq refined the prompt and Hugging Face generated the image artifact.",
            "prompt": refined_prompt,
        }

    def explanation_visual(
        self,
        prompt: str,
        diagram_type: str = "concept_map",
        theme: str = "modern-3d-education",
    ) -> Dict[str, Any]:
        messages = [
            ChatMessage(
                role="system",
                content=(
                    "You are a visualization architect. Return strict JSON only. "
                    "Build a valid diagram spec with this shape: "
                    '{"type":"diagram","diagramType":"...","title":"...","nodes":[...],"edges":[...],"style":{...}}. '
                    "Use 5-9 nodes, meaningful labels, and relationships that explain the topic clearly."
                ),
            ),
            ChatMessage(
                role="user",
                content=(
                    f"Topic: {prompt}\n"
                    f"diagramType: {diagram_type}\n"
                    f"theme: {theme}\n"
                    "Return JSON only."
                ),
            ),
        ]

        raw = self.clients.groq_chat(
            messages,
            model=settings.groq_model_quality,
            temperature=0.1,
        )

        parsed: Dict[str, Any] | None = None
        try:
            parsed = json.loads(raw.strip())
        except Exception:
            match = re.search(r"\{[\s\S]*\}", raw)
            if match:
                try:
                    parsed = json.loads(match.group(0))
                except Exception:
                    parsed = None

        if not isinstance(parsed, dict):
            parsed = {
                "type": "diagram",
                "diagramType": diagram_type,
                "title": "Concept Overview",
                "subtitle": prompt[:100],
                "nodes": [
                    {
                        "id": "n1",
                        "type": "concept",
                        "label": "Problem",
                        "description": "Core topic",
                    },
                    {
                        "id": "n2",
                        "type": "process",
                        "label": "Approach",
                        "description": "How it works",
                    },
                    {
                        "id": "n3",
                        "type": "output",
                        "label": "Outcome",
                        "description": "Expected result",
                    },
                ],
                "edges": [
                    {
                        "source": "n1",
                        "target": "n2",
                        "label": "informs",
                        "type": "smoothstep",
                        "animated": True,
                    },
                    {
                        "source": "n2",
                        "target": "n3",
                        "label": "produces",
                        "type": "smoothstep",
                        "animated": True,
                    },
                ],
                "style": {
                    "theme": theme,
                    "layout": "hierarchical",
                    "edgeStyle": "gradient",
                    "nodeShape": "card",
                    "animation": "subtle",
                },
            }

        parsed["type"] = "diagram"
        parsed.setdefault("diagramType", diagram_type)
        parsed.setdefault("style", {})
        if isinstance(parsed["style"], dict):
            parsed["style"].setdefault("theme", theme)

        visual_block = (
            '<visual type="diagram">\n'
            + json.dumps(parsed, ensure_ascii=True, indent=2)
            + "\n</visual>"
        )
        return {
            "status": "ok",
            "type": "diagram",
            "diagram": parsed,
            "visual_block": visual_block,
        }
