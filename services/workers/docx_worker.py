from pathlib import Path
from docx import Document
from datetime import datetime


class DOCXWorker:
    """Generate DOCX artifacts."""
    
    def __init__(self, output_dir: str = "data/artifacts"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate(self, content: str, title: str = "Document") -> str:
        """Generate DOCX from text content."""
        
        # Sanitize filename
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_title}_{timestamp}.docx"
        filepath = self.output_dir / filename
        
        # Create document
        doc = Document()
        doc.add_heading(title, 0)
        
        # Add content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        doc.save(str(filepath))
        
        return str(filepath)
